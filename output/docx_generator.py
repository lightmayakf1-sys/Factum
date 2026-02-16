"""Генерация DOCX-карточки оборудования (python-docx)."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from gigachat_api.schema import ExtractedValue, CHECKLIST_FIELDS, SECTION_GROUPS
from output.canonical import source_display, missing_param_note
from output.formatter import format_value


def generate_card(
    resolved: dict[str, ExtractedValue | None],
    notes: list[str],
    overview: str = "",
    output_path: Path | None = None,
) -> Document:
    """Сгенерировать DOCX-карточку оборудования.

    Args:
        resolved: Финальные данные (field_name -> ExtractedValue).
        notes: Список примечаний.
        overview: Инженерный обзор (3-5 строк).
        output_path: Путь для сохранения. Если None — не сохраняет.

    Returns:
        Document: объект python-docx.
    """
    doc = Document()

    # Стили
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(10)

    # --- Заголовок ---
    model = _get_value(resolved, "a2_model", "—")
    manufacturer = _get_value(resolved, "a3_manufacturer", "—")

    title = doc.add_heading(level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"КАРТОЧКА ОБОРУДОВАНИЯ: {model} — {manufacturer}")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"Ревизия: 1 • Дата формирования: {date.today().strftime('%d.%m.%Y')}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 100, 100)

    # --- Инженерный обзор ---
    if overview:
        doc.add_heading("Инженерный обзор", level=2)
        doc.add_paragraph(overview)

    # --- Группы параметров ---
    for group_key in ["A", "B", "C", "D", "E", "F", "G", "H"]:
        group_title, field_names = SECTION_GROUPS[group_key]
        doc.add_heading(group_title, level=2)

        if group_key == "A":
            # Группа A без столбца "Источник"
            table = doc.add_table(rows=1, cols=2)
        else:
            # Группы B–H с тремя столбцами
            table = doc.add_table(rows=1, cols=3)

        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Заголовки
        header_cells = table.rows[0].cells
        _set_cell(header_cells[0], "Параметр", bold=True)
        _set_cell(header_cells[1], "Значение", bold=True)
        if group_key != "A":
            _set_cell(header_cells[2], "Источник / Статус", bold=True)

        # Данные
        for field_name in field_names:
            label = dict(CHECKLIST_FIELDS).get(field_name, field_name)
            ev = resolved.get(field_name)

            row = table.add_row()
            _set_cell(row.cells[0], label)

            if ev is None:
                _set_cell(row.cells[1], "нет данных", italic=True)
                if group_key != "A":
                    _set_cell(row.cells[2], "—")
            else:
                display_value = format_value(ev.value)

                if ev.status and "конфликт" not in ev.status.lower():
                    display_value = f"{display_value} [{ev.status}]"

                if ev.status and "конфликт" in ev.status.lower():
                    # Конфликт — показываем выбранное значение зелёным + метку
                    _set_cell_conflict_value(row.cells[1], ev)
                else:
                    _set_cell(row.cells[1], display_value)

                if group_key != "A":
                    src = source_display(
                        file=ev.source.file,
                        doc_type=ev.source.doc_type,
                        page=ev.source.page,
                        section=ev.source.section,
                        quote=ev.source.quote,
                        confidence=ev.source.confidence,
                    )
                    _set_cell(row.cells[2], src,
                              color=RGBColor(180, 0, 0) if ev.source.confidence == "low" else None)

        # Ширина столбцов
        if group_key == "A":
            _set_col_width(table, 0, Cm(6))
            _set_col_width(table, 1, Cm(11))
        else:
            _set_col_width(table, 0, Cm(5))
            _set_col_width(table, 1, Cm(5))
            _set_col_width(table, 2, Cm(7))

    # --- ПРИМЕЧАНИЯ ---
    _write_notes_section(doc, resolved, notes)

    # Сохранение
    if output_path:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))

    return doc


def _set_cell_conflict_value(cell, ev: ExtractedValue):
    """Записать конфликтное значение в ячейку с выделением."""
    cell.text = ""
    paragraph = cell.paragraphs[0]

    # Заголовок: выбранное значение зелёным
    run = paragraph.add_run(f"✔ {format_value(ev.value)}")
    run.font.size = Pt(9)
    run.font.name = "Times New Roman"
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 128, 0)  # Зелёный

    run = paragraph.add_run("  [конфликт]")
    run.font.size = Pt(8)
    run.font.name = "Times New Roman"
    run.font.color.rgb = RGBColor(200, 120, 0)  # Оранжевый

    # Остальные варианты серым
    if ev.conflict_values:
        for entry in ev.conflict_values:
            if entry.is_selected:
                continue
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(f"✖ {format_value(entry.value)}")
            run.font.size = Pt(8)
            run.font.name = "Times New Roman"
            run.font.color.rgb = RGBColor(130, 130, 130)  # Серый

            src_text = f"  ({entry.source.file}"
            if entry.source.doc_type:
                src_text += f", {entry.source.doc_type}"
            if entry.source.page:
                src_text += f", стр. {entry.source.page}"
            src_text += ")"
            run = p.add_run(src_text)
            run.font.size = Pt(7)
            run.font.name = "Times New Roman"
            run.font.color.rgb = RGBColor(160, 160, 160)


def _write_notes_section(doc: Document, resolved: dict, extra_notes: list[str]):
    """Секция ПРИМЕЧАНИЯ с красивым отображением конфликтов."""
    # Собираем всё
    conflict_notes = []
    other_notes = []

    # 1. Конфликты — в структурированном виде
    for field_name, label in CHECKLIST_FIELDS:
        ev = resolved.get(field_name)
        if ev and ev.status and "конфликт" in ev.status.lower():
            conflict_notes.append((label, ev))

    # 2. Низкая уверенность
    for field_name, label in CHECKLIST_FIELDS:
        ev = resolved.get(field_name)
        if ev and ev.source.confidence == "low":
            other_notes.append(f"{label} — считан с низким качеством, требует проверки.")

    # 3. Пропуски
    for field_name, label in CHECKLIST_FIELDS:
        ev = resolved.get(field_name)
        if ev is None:
            other_notes.append(missing_param_note(label))

    # 4. Дополнительные примечания из верификации
    for note in extra_notes:
        if note not in other_notes:
            other_notes.append(note)

    has_content = conflict_notes or other_notes
    if not has_content:
        return

    doc.add_heading("ПРИМЕЧАНИЯ", level=2)
    note_num = 1

    # --- Конфликты: отдельная таблица для каждого ---
    if conflict_notes:
        p = doc.add_paragraph()
        run = p.add_run("⚠ РАСХОЖДЕНИЯ МЕЖДУ ИСТОЧНИКАМИ")
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor(200, 120, 0)

        for label, ev in conflict_notes:
            p = doc.add_paragraph()
            run = p.add_run(f"{note_num}. {label}")
            run.font.size = Pt(9)
            run.font.bold = True
            run.font.name = "Times New Roman"
            note_num += 1

            if ev.conflict_values:
                # Таблица конфликтов
                tbl = doc.add_table(rows=1, cols=3)
                tbl.style = "Table Grid"

                hdr = tbl.rows[0].cells
                _set_cell(hdr[0], "№", bold=True)
                _set_cell(hdr[1], "Значение", bold=True)
                _set_cell(hdr[2], "Источник", bold=True)

                for idx, entry in enumerate(ev.conflict_values, 1):
                    row = tbl.add_row()
                    marker = "✔" if entry.is_selected else str(idx)

                    if entry.is_selected:
                        _set_cell(row.cells[0], marker, bold=True,
                                  color=RGBColor(0, 128, 0))
                        _set_cell(row.cells[1], format_value(entry.value), bold=True,
                                  color=RGBColor(0, 128, 0))
                    else:
                        _set_cell(row.cells[0], marker)
                        _set_cell(row.cells[1], format_value(entry.value),
                                  color=RGBColor(130, 130, 130))

                    src = source_display(
                        file=entry.source.file,
                        doc_type=entry.source.doc_type,
                        page=entry.source.page,
                        section=entry.source.section,
                        quote=entry.source.quote,
                        confidence=entry.source.confidence,
                    )
                    _set_cell(row.cells[2], src)

                _set_col_width(tbl, 0, Cm(1))
                _set_col_width(tbl, 1, Cm(7))
                _set_col_width(tbl, 2, Cm(9))
            else:
                # Нет структурированных данных — обычный текст
                p = doc.add_paragraph(ev.note or "Расхождение между источниками.")
                run = p.runs[0] if p.runs else p.add_run("")
                run.font.size = Pt(9)

        doc.add_paragraph()  # Отступ

    # --- Остальные примечания ---
    for note in other_notes:
        doc.add_paragraph(f"{note_num}. {note}")
        note_num += 1


def _get_value(resolved: dict, field: str, default: str = "") -> str:
    ev = resolved.get(field)
    if ev is None:
        return default
    return ev.value or default


def _set_cell(cell, text: str, bold: bool = False, italic: bool = False,
              color: RGBColor | None = None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.font.size = Pt(9)
    run.font.name = "Times New Roman"
    if bold:
        run.font.bold = True
    if italic:
        run.font.italic = True
    if color:
        run.font.color.rgb = color


def _set_col_width(table, col_idx: int, width):
    for row in table.rows:
        row.cells[col_idx].width = width
